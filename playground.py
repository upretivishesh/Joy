import json
import os
import re
import smtplib
from collections import Counter
from datetime import datetime
from email.message import EmailMessage
from email.utils import formataddr
from io import BytesIO
from pathlib import Path
from pdf2image import convert_from_bytes
import pytesseract
from concurrent.futures import ThreadPoolExecutor, as_completed


import pandas as pd
import streamlit as st

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from docx import Document
except Exception:
    Document = None


APP_NAME = "Joy"
DEFAULT_COMPANY = "Seven Hiring"
DATA_DIR = Path(".joy_data")

KNOWN_SKILLS = [
    "python", "java", "javascript", "typescript", "react", "node", "django",
    "flask", "fastapi", "sql", "postgresql", "mysql", "mongodb", "aws",
    "azure", "gcp", "docker", "kubernetes", "terraform", "linux", "git",
    "machine learning", "data analysis", "excel", "power bi", "tableau",
    "salesforce", "sap", "erp", "crm", "leadership", "communication",
    "project management", "agrochemical", "pharma", "fmcg", "sales",
    "distribution", "channel sales", "key account", "hplc", "gc-ms", "gmp",
    "qa", "qc", "production", "manufacturing", "finance", "accounting",
]

STOP_WORDS = {
    "about", "above", "after", "again", "against", "also", "among", "and",
    "any", "are", "based", "been", "being", "best", "between", "both",
    "business", "candidate", "candidates", "client", "company", "description",
    "detail", "details", "development", "each", "etc", "experience",
    "experienced", "for", "from", "good", "have", "hiring", "including",
    "india", "into", "job", "knowledge", "like", "looking", "management",
    "manager", "must", "need", "needs", "our", "over", "preferred", "profile",
    "required", "requirements", "responsibilities", "role", "should", "skill",
    "skills", "strong", "team", "that", "the", "their", "this", "through",
    "with", "work", "working", "years", "you", "your",
}

DEFAULT_QUESTIONS = [
    "Current CTC",
    "Expected CTC",
    "Notice period",
    "Current location",
    "Preferred work location",
    "Total experience",
    "Reason for job change",
    "Current company and designation",
    "Any offer in hand",
    "Suitable slot for a 5-minute discussion",
]


def init_state() -> None:
    defaults = {
        "gmail_authenticated": False,
        "results_df": pd.DataFrame(),
        "last_role": "",
        "last_jd": "",
        "last_keywords": [],
        "email_results": [],
        "questions_text": "\n".join(DEFAULT_QUESTIONS),
        "sender_email": "",
        "sender_password": "",
        "sender_name": "",
        "company_name": DEFAULT_COMPANY,
        "upload_session": 0,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def get_secret(name: str, default: str = "") -> str:
    try:
        value = st.secrets.get(name, "")
    except Exception:
        value = ""
    return value or os.getenv(name, default)


def normalize_whitespace(text: str) -> str:
    text = text or ""
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def name_from_email_address(email: str) -> str:
    local = (email or "").split("@")[0]
    local = re.sub(r"[^A-Za-z]+", " ", local)
    parts = [part for part in local.split() if len(part) > 1 and not part.isdigit()]
    if not parts:
        return ""
    return " ".join(part.capitalize() for part in parts[:3])


def mask_email(email: str) -> str:
    if "@" not in email:
        return email
    local, domain = email.split("@", 1)
    if len(local) <= 2:
        masked_local = local[:1] + "*"
    else:
        masked_local = local[:2] + "*" * min(5, len(local) - 2)
    return f"{masked_local}@{domain}"


def login_user(email: str, app_password: str, sender_name: str, company_name: str) -> None:
    clean_email = email.strip().lower()
    st.session_state.gmail_authenticated = True
    st.session_state.sender_email = clean_email
    st.session_state.sender_password = app_password.replace(" ", "").strip()
    st.session_state.sender_name = sender_name.strip() or name_from_email_address(clean_email)
    st.session_state.company_name = company_name.strip() or DEFAULT_COMPANY


def logout_user() -> None:
    for key in [
    "gmail_authenticated",
    "sender_email",
    "sender_password",
    "sender_name",
    "company_name",
]:
        st.session_state[key] = False if key == "gmail_authenticated" else ""
    st.session_state.email_results = []
    st.rerun()


def safe_filename_part(value: str) -> str:
    clean = re.sub(r"[^a-zA-Z0-9._-]+", "_", value or "user")
    return clean.strip("_")[:80] or "user"


def history_path(user_key: str) -> Path:
    DATA_DIR.mkdir(exist_ok=True)
    return DATA_DIR / f"history_{safe_filename_part(user_key)}.csv"

@st.cache_data(show_spinner=False)
def read_uploaded_file(file_name, data) -> tuple[str, str]:
    name = file_name.lower()

    try:
        if name.endswith(".pdf"):

            if pdfplumber is None:
                return "", "pdfplumber is not installed."

            with pdfplumber.open(BytesIO(data)) as pdf:

                pages = pdf.pages[:5]

                text = "\n".join(
                    page.extract_text() or ""
                    for page in pages
                )

            text = text.strip()

            # OCR fallback for scanned/image PDFs
            if len(text) < 35:

                ocr_text = ocr_pdf(data)

                if len(ocr_text) > len(text):
                    text = ocr_text

            return text.strip(), ""

        if name.endswith(".txt"):
            return data.decode("utf-8", errors="ignore").strip(), ""
        
        if name.endswith(".docx"):
            if Document is None:
                return "", "python-docx is not installed."

            try:
                doc = Document(BytesIO(data))

                text = "\n".join(p.text for p in doc.paragraphs)

                if not text.strip():
                    tables_text = []

                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " ".join(cell.text for cell in row.cells)
                            tables_text.append(row_text)

                    text = "\n".join(tables_text)

                if not text.strip():
                    return "", "DOCX opened but no readable text found."

                return text.strip(), ""

            except Exception as exc:
                return "", f"DOCX parsing failed: {exc}"

        return "", "Unsupported file type."
    except Exception as exc:
        return "", f"Could not read file: {exc}"

def ocr_pdf(data: bytes) -> str:

    try:

        images = convert_from_bytes(
            data,
            dpi=170,
            first_page=1,
            last_page=2,
            thread_count=1,
            grayscale=True,
        )

        text_parts = []

        for image in images:

            text = pytesseract.image_to_string(image)

            if text.strip():
                text_parts.append(text)

        return "\n".join(text_parts).strip()

    except Exception:
        return ""

def extract_email(text: str) -> str:
    match = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    patterns = [
        r"(?:\+91[-.\s]?)?[6-9]\d{9}",
        r"\+?\d{1,4}[-.\s]?\(?\d{2,5}\)?[-.\s]?\d{3,5}[-.\s]?\d{3,6}",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    return ""


def clean_name_candidate(value: str) -> str:
    value = re.sub(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", " ", value)
    value = re.sub(r"(?:\+91[-.\s]?)?[6-9]\d{9}", " ", value)
    value = re.sub(r"https?://\S+|www\.\S+", " ", value, flags=re.I)
    value = re.sub(r"[^A-Za-z .'-]", " ", value)
    value = normalize_whitespace(value)
    if value.isupper():
        value = value.title()
    return value.strip(" .'-")


def extract_name_from_email(text: str) -> str:
    email = extract_email(text)
    if not email:
        return ""
    name = name_from_email_address(email)
    banned = {"info", "admin", "resume", "career", "careers", "mail", "email"}
    parts = [part for part in name.split() if part.lower() not in banned]
    return " ".join(parts) if len(parts) >= 2 else ""


def extract_name(text: str, filename: str = "") -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    email_name = extract_name_from_email(text)
    email_tokens = {token.lower() for token in email_name.split()}
    bad_phrases = {
        "resume", "curriculum", "vitae", "profile", "email", "phone", "mobile",
        "linkedin", "github", "portfolio", "address", "summary", "objective",
        "education", "experience", "employment", "project", "projects", "skills",
        "certification", "certifications",

        # job titles
        "developer", "engineer", "manager", "analyst", "consultant",
        "specialist", "executive", "assistant", "designer", "architect",
        "lead", "intern", "recruiter", "marketer", "accountant",
        "graphic designer", "ui ux", "ui/ux", "sales executive",
        "hr executive", "software engineer", "data analyst",

        # random resume junk
        "curriculum vitae", "professional summary",
        "work experience", "career objective",
    }

    candidates: list[tuple[int, str]] = []
    for idx, line in enumerate(lines[:28]):
        lower_line = line.lower()
        if "@" in line or "http" in lower_line or ":" in line:
            continue

        clean = clean_name_candidate(line)
        if not clean:
            continue

        lower = clean.lower()
        words = clean.split()
        if not 2 <= len(words) <= 5:
            continue
        if any(char.isdigit() for char in clean):
            continue
        if any(
            phrase == lower or phrase in lower.split()
            for phrase in bad_phrases
        ):
            continue
        if len(clean) > 55:
            continue
        if not all(re.fullmatch(r"[A-Za-z][A-Za-z.'-]*", word) for word in words):
            continue

        score = 80 - idx * 2
        if 2 <= len(words) <= 3:
            score += 22
        if len(words) == 1:
            score -= 20
        if email_tokens and any(word.lower() in email_tokens for word in words):
            score += 35
        if all(word[:1].isupper() for word in words):
            score += 12
        if idx <= 3:
            score += 12
        if idx == 0:
            score += 25
        if idx == 1:
            score += 18    
        candidates.append((score, clean.title() if clean.isupper() else clean))

    if candidates:
        candidates.sort(reverse=True, key=lambda item: item[0])
        if candidates[0][0] >= 62:
            return candidates[0][1]

    if email_name:
        return email_name

    return "Unknown Candidate"


def extract_experience(text: str) -> float:

    text = text.lower()

    # explicit experience mentions first
    explicit_patterns = [

        r"(\d{1,2}(?:\.\d+)?)\s*\+?\s*(?:years|year|yrs|yr)",

        r"experience\s*[:\-]?\s*(\d{1,2}(?:\.\d+)?)",

        r"total\s*experience\s*[:\-]?\s*(\d{1,2}(?:\.\d+)?)",

        r"overall\s*experience\s*[:\-]?\s*(\d{1,2}(?:\.\d+)?)",

        r"professional\s*experience\s*[:\-]?\s*(\d{1,2}(?:\.\d+)?)",

        r"worked\s*for\s*(\d{1,2}(?:\.\d+)?)\s*(?:years|yrs)",

    ]

    for pattern in explicit_patterns:
        match = re.search(pattern, text, flags=re.I)
        if match:
            try:
                years = float(match.group(1))

                if 0 <= years <= 50:
                    return round(years, 1)

            except:
                pass

    # fallback → calculate from years in resume
    years = re.findall(r"(20\d{2}|19\d{2})", text)

    years = [int(y) for y in years]

    if len(years) >= 2:

        min_year = min(years)
        max_year = max(years)

        current_year = datetime.now().year

        if max_year > current_year:
            max_year = current_year

        exp = max_year - min_year

        if 0 <= exp <= 40:
            return float(exp)

    return 0.0


def extract_skills(text: str) -> list[str]:
    lower = text.lower()
    return sorted({skill for skill in KNOWN_SKILLS if skill in lower})


def clean_role_title(value: str) -> str:
    value = normalize_whitespace(value)
    value = re.split(
        r"\b(location|experience|department|reports|reporting|salary|ctc|about|overview|responsibilities)\b",
        value,
        flags=re.I,
    )[0]
    value = re.sub(r"[^A-Za-z0-9 /&+.,'-]", " ", value)
    value = normalize_whitespace(value).strip(" -:.,")
    value = re.sub(r"^(for|as|a|an|the)\s+", "", value, flags=re.I)
    words = value.split()
    if len(words) > 8:
        value = " ".join(words[:8])
    return value.title() if value else ""


def extract_role_from_jd(jd_text: str, fallback: str = "") -> str:

    if fallback.strip():
        return clean_role_title(fallback)

    text = jd_text or ""

    if not text.strip():
        return "Open Role"

    lines = [
        normalize_whitespace(line)
        for line in text.splitlines()
        if normalize_whitespace(line)
    ]

    # common actual job titles
    title_keywords = {
        "engineer", "developer", "designer", "manager", "executive",
        "specialist", "associate", "lead", "architect", "consultant",
        "analyst", "officer", "head", "director", "intern",
        "coordinator", "writer", "recruiter", "marketer",
        "sales", "hr", "product", "ui", "ux", "qa",
        "accountant", "teacher", "trainer", "chef"
    }

    # garbage phrases
    bad_phrases = {
        "about us", "job description", "responsibilities",
        "requirements", "qualifications", "preferred",
        "mandatory", "salary", "location", "experience",
        "apply now", "about company", "skills required",
        "candidate profile", "company overview"
    }

    patterns = [
        r"(?:job\s*)?title\s*[:\-]\s*(.+)",
        r"(?:role|position|designation)\s*[:\-]\s*(.+)",
        r"hiring\s+(?:for\s+)?(.+)",
        r"opening\s+(?:for\s+)?(.+)",
        r"urgent\s+(?:requirement\s+for\s+)?(.+)",
        r"looking\s+for\s+(.+)",
    ]

    scored_candidates = []

    # ---------- regex extraction ----------
    for idx, line in enumerate(lines[:40]):

        lower = line.lower()

        if any(bad in lower for bad in bad_phrases):
            continue

        for pattern in patterns:

            match = re.search(pattern, line, flags=re.I)

            if match:

                role = clean_role_title(match.group(1))

                if not role:
                    continue

                score = 100 - idx

                if len(role.split()) <= 8:
                    score += 10

                scored_candidates.append((score, role))

    # ---------- heading detection ----------
    for idx, line in enumerate(lines[:15]):

        clean = clean_role_title(line)

        if not clean:
            continue

        lower = clean.lower()

        if any(bad in lower for bad in bad_phrases):
            continue

        words = clean.split()

        # skip giant paragraphs
        if len(words) > 8:
            continue

        score = 60 - idx

        # title keyword bonus
        if any(
            keyword in lower
            for keyword in title_keywords
        ):
            score += 40

        # uppercase headings bonus
        if line.isupper():
            score += 15

        # top-of-document bonus
        if idx <= 3:
            score += 20

        # penalize weird lines
        if ":" in line:
            score -= 10

        scored_candidates.append((score, clean))

    # ---------- final selection ----------
    if scored_candidates:

        scored_candidates.sort(
            key=lambda x: x[0],
            reverse=True
        )

        best_role = scored_candidates[0][1]

        return clean_role_title(best_role)

    return "Open Role"


def ai_extract_role_title(jd_text: str, api_key: str, model: str) -> str:
    if not api_key or not jd_text.strip():
        return ""
    try:
        from openai import OpenAI

        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": "Extract the exact hiring role title from a job description. Return JSON only.",
                },
                {
                    "role": "user",
                    "content": f"Return only JSON like {{\"title\":\"\"}}. JD:\n{jd_text[:3500]}",
                },
            ],
            temperature=0,
            max_tokens=80,
        )
        raw = re.sub(r"```json|```", "", response.choices[0].message.content or "").strip()
        title = clean_role_title(json.loads(raw).get("title", ""))
        return title if title and title != "Open Role" else ""
    except Exception:
        return ""


def detect_role_title(jd_text: str, fallback: str, api_key: str, model: str) -> str:
    if fallback.strip():
        return clean_role_title(fallback)
    ai_title = ai_extract_role_title(jd_text, api_key, model)
    return ai_title or extract_role_from_jd(jd_text)


def parse_min_experience(jd_text: str) -> float:
    patterns = [
        r"(\d{1,2})\s*[-+]\s*\d{0,2}\s*(?:years?|yrs?)",
        r"minimum\s*(?:of)?\s*(\d{1,2})\s*(?:years?|yrs?)",
        r"at least\s*(\d{1,2})\s*(?:years?|yrs?)",
        r"(\d{1,2})\+?\s*(?:years?|yrs?)\s*(?:of)?\s*experience",
    ]
    for pattern in patterns:
        match = re.search(pattern, jd_text, flags=re.I)
        if match:
            try:
                return float(match.group(1))
            except ValueError:
                return 0.0
    return 0.0


def extract_keywords(text: str, extra_keywords: str = "", limit: int = 35) -> list[str]:
    text = text or ""
    configured = [kw.strip().lower() for kw in extra_keywords.split(",") if kw.strip()]
    skill_hits = [skill for skill in KNOWN_SKILLS if skill in text.lower()]

    words = re.findall(r"\b[a-zA-Z][a-zA-Z+#.-]{2,}\b", text.lower())
    words = [w.strip(".-") for w in words if w not in STOP_WORDS and len(w) >= 3]
    common = [word for word, _ in Counter(words).most_common(limit)]

    keywords = []
    for item in configured + skill_hits + common:
        if item and item not in keywords and item not in STOP_WORDS:
            keywords.append(item)
    return keywords[:limit]


def keyword_match_score(resume_text: str, keywords: list[str]) -> tuple[int, list[str], list[str]]:
    if not keywords:
        return 50, [], []

    lower = resume_text.lower()
    matched = [kw for kw in keywords if kw.lower() in lower]
    missing = [kw for kw in keywords if kw.lower() not in lower]
    score = round((len(matched) / len(keywords)) * 100)
    return int(score), matched, missing


def experience_score(candidate_years: float, required_years: float) -> int:
    if required_years <= 0:
        return 70 if candidate_years == 0 else 85
    if candidate_years >= required_years:
        return 100
    return int(max(0, min(100, (candidate_years / required_years) * 100)))


def contact_score(email: str, phone: str) -> int:
    score = 0
    if email:
        score += 60
    if phone:
        score += 40
    return score


def ai_score_resume(
    jd_text: str,
    resume_text: str,
    role: str,
    api_key: str,
    model: str,
) -> tuple[int | None, str]:
    if not api_key:
        return None, ""

    try:
        from openai import OpenAI

        client = OpenAI(api_key=api_key)
        prompt = f"""
Score the resume against the role and job description.
Return JSON only:
{{"score": 0, "reason": ""}}

Role: {role}

Job description:
{jd_text[:2500]}

Resume:
{resume_text[:3500]}
"""
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": "You are a strict recruiter. Reward direct evidence, penalize missing must-haves, and stay concise.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0,
            max_tokens=120,
            timeout=20,
        )
        raw = response.choices[0].message.content or "{}"
        raw = re.sub(r"```json|```", "", raw).strip()
        data = json.loads(raw)
        score = int(float(data.get("score", 0)))
        reason = str(data.get("reason", "")).strip()
        return max(0, min(100, score)), reason
    except Exception as exc:
        return None, f"AI scoring skipped: {exc}"


def make_reason(matched: list[str], missing: list[str], exp: float, min_exp: float) -> str:
    matched_text = ", ".join(matched[:5]) if matched else "few direct keyword matches"
    missing_text = ", ".join(missing[:4]) if missing else "no obvious must-have gaps"
    if min_exp > 0:
        exp_text = f"{exp:g} yrs found vs {min_exp:g}+ yrs expected"
    else:
        exp_text = f"{exp:g} yrs found" if exp else "experience not clearly stated"
    return f"Matched {matched_text}. Missing/unclear: {missing_text}. {exp_text}."


def verdict_from_score(score: float) -> str:
    if score >= 82:
        return "Strong Fit"
    if score >= 68:
        return "Good Fit"
    if score >= 50:
        return "Review"
    return "Low Fit"


def score_resume(
    jd_text: str,
    role: str,
    resume_text: str,
    filename: str,
    keywords: list[str],
    min_exp: float,
    api_key: str,
    model: str,
) -> dict:
    name = extract_name(resume_text, filename)
    email = extract_email(resume_text)
    phone = extract_phone(resume_text)
    exp = extract_experience(resume_text)
    skills = extract_skills(resume_text)

    kw_score, matched, missing = keyword_match_score(resume_text, keywords)
    exp_score = experience_score(exp, min_exp)
    cnt_score = contact_score(email, phone)
    skill_score = min(100, len(skills) * 12)

    heuristic = (kw_score * 0.55) + (exp_score * 0.2) + (skill_score * 0.15) + (cnt_score * 0.1)
    ai_score = None
    ai_reason = ""

    # only AI-score promising resumes
    if kw_score >= 35 or exp_score >= 55:

        ai_score, ai_reason = ai_score_resume(
            jd_text,
            resume_text,
            role,
            api_key,
            model,
        )

    if ai_score is None:
        final_score = round(heuristic, 1)
        reason = ai_reason or make_reason(matched, missing, exp, min_exp)
        ai_used = False
    else:
        final_score = round((heuristic * 0.45) + (ai_score * 0.55), 1)
        reason = ai_reason or make_reason(matched, missing, exp, min_exp)
        ai_used = True

    return {
        "Send": verdict_from_score(final_score) in {"Strong Fit", "Good Fit"} and bool(email),
        "Name": name,
        "Email": email,
        "Phone": phone,
        "Experience": exp,
        "Keyword Score": kw_score,
        "Final Score": final_score,
        "Verdict": verdict_from_score(final_score),
        "Matched Keywords": ", ".join(matched[:12]),
        "Missing Keywords": ", ".join(missing[:10]),
        "Skills": ", ".join(skills[:12]),
        "Reason": reason,
        "Source File": filename,
        "AI Used": ai_used,
    }


def save_history(df: pd.DataFrame, role: str, user_key: str) -> None:
    if df.empty:
        return
    path = history_path(user_key)
    batch = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    to_save = df.copy()
    to_save["Role"] = role
    to_save["Screened At"] = batch

    if path.exists():
        old = pd.read_csv(path)
        to_save = pd.concat([old, to_save], ignore_index=True)
        to_save = to_save.drop_duplicates(
            subset=["Email", "Role", "Source File"],
            keep="last"
        )
    to_save.to_csv(path, index=False)


def load_history(user_key: str) -> pd.DataFrame:
    path = history_path(user_key)
    if path.exists():
        return pd.read_csv(path)
    return pd.DataFrame()


def clear_history(user_key: str) -> None:
    path = history_path(user_key)
    if path.exists():
        path.unlink()

def clear_role_history(user_key: str, role: str) -> None:

    path = history_path(user_key)

    if not path.exists():
        return

    df = pd.read_csv(path)

    if "Role" not in df.columns:
        return

    df = df[df["Role"] != role]

    df.to_csv(path, index=False)        

def reset_screening_session() -> None:
    st.session_state.results_df = pd.DataFrame()
    st.session_state.email_results = []
    st.session_state.last_role = ""
    st.session_state.last_jd = ""
    st.session_state.last_keywords = []

    # reset text areas
    st.session_state["typed_jd_text"] = ""
    st.session_state["role_input"] = ""
    st.session_state["extra_keywords"] = ""

    # force uploader refresh
    st.session_state.upload_session += 1

    st.session_state["_reset_done"] = True

def questions_from_text(text: str) -> list[str]:
    questions = []
    for line in text.splitlines():
        clean = re.sub(r"^\s*[-*0-9.)]+\s*", "", line).strip()
        if clean:
            questions.append(clean)
    return questions or DEFAULT_QUESTIONS


def first_name(full_name: str) -> str:
    if not full_name or full_name == "Unknown Candidate":
        return "there"
    return full_name.split()[0].strip(",")


def build_email_body(
    candidate: pd.Series,
    role: str,
    sender_name: str,
    company_name: str,
    questions: list[str],
    extra_note: str,
) -> str:
    greeting = first_name(str(candidate.get("Name", "")))
    company = company_name or DEFAULT_COMPANY
    score = candidate.get("Final Score", "")
    verdict = candidate.get("Verdict", "")

    lines = [
        "Hi {first_name},",
        "",
        f"I reviewed your profile for the {role} role and it looks relevant for the first screening round.",
        "",
    ]

    if extra_note.strip():
        lines.extend([extra_note.strip(), ""])

    lines.extend(
        [
            "To move ahead without a back-and-forth call, please reply with these details:",
            "",
        ]
    )

    for idx, question in enumerate(questions, start=1):
        lines.append(f"{idx}. {question}")

    lines.extend(
        [
            "",
            "Once I have this, I can confirm fit, share the next step, and avoid asking you the same basics again on call.",
            "",
            "Best regards,",
            sender_name or "Recruitment Team",
            company,
        ]
    )
    return "\n".join(lines)


def send_email(
    sender_email: str,
    sender_password: str,
    sender_name: str,
    recipient_email: str,
    subject: str,
    body: str,
) -> tuple[bool, str]:
    try:
        msg = EmailMessage()
        msg["From"] = formataddr((sender_name or sender_email, sender_email))
        msg["To"] = recipient_email
        msg["Subject"] = subject
        safe_body = (
            body.encode("utf-8", errors="ignore")
            .decode("utf-8", errors="ignore")
            .replace(chr(10), "<br>")
        )

        html_body = f"""
        <html>
        <body style="font-family:Aptos, Arial, sans-serif; font-size:12pt; line-height:1.6;">
        {safe_body}
        </body>
        </html>
        """

        msg.set_content(body)

        msg.add_alternative(html_body, subtype="html")

        with smtplib.SMTP("smtp.gmail.com", 587, timeout=30) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
        return True, "Sent"
    except smtplib.SMTPAuthenticationError:
        return False, "Gmail rejected the login. Use a Gmail App Password, not your normal password."
    except Exception as exc:
        print("EMAIL ERROR:", exc)
        return False, str(exc)

def render_template_variables(text: str, candidate: pd.Series, role: str) -> str:

    variables = {
        "{first_name}": first_name(str(candidate.get("Name", ""))),
        "{full_name}": str(candidate.get("Name", "")),
        "{role}": role,
        "{email}": str(candidate.get("Email", "")),
        "{phone}": str(candidate.get("Phone", "")),
        "{experience}": str(candidate.get("Experience", "")),
        "{score}": str(candidate.get("Final Score", "")),
        "{verdict}": str(candidate.get("Verdict", "")),
    }

    rendered = text

    for key, value in variables.items():
        rendered = rendered.replace(key, value)

    return rendered

def send_bulk_emails(
    selected_df: pd.DataFrame,
    role: str,
    sender_email: str,
    sender_password: str,
    sender_name: str,
    company_name: str,
    subject: str,
    questions: list[str],
    extra_note: str,
    custom_body: str = "",
) -> list[dict]:
    results = []


    def email_worker(candidate):

        name = str(candidate.get("Name", "Candidate"))

        recipient = str(candidate.get("Email", "")).strip()

        if "@" not in recipient:

            return {
                "Name": name,
                "Email": recipient,
                "Success": False,
                "Message": "Missing email",
            }

        if custom_body.strip():

            body = render_template_variables(
                custom_body,
                candidate,
                role,
            )

        else:

            body = build_email_body(
                candidate,
                role,
                sender_name,
                company_name,
                questions,
                extra_note,
            )

        personalized_subject = render_template_variables(
            subject,
            candidate,
            role,
        )

        ok, message = send_email(
            sender_email,
            sender_password,
            sender_name,
            recipient,
            personalized_subject,
            body,
        )

        return {
            "Name": name,
            "Email": recipient,
            "Success": ok,
            "Message": message,
        }

    with ThreadPoolExecutor(max_workers=2) as executor:

        futures = [
            executor.submit(email_worker, candidate)
            for _, candidate in selected_df.iterrows()
        ]

        for future in as_completed(futures):

            try:
                results.append(future.result())

            except Exception as exc:

                results.append({
                    "Name": "Unknown",
                    "Email": "",
                    "Success": False,
                    "Message": str(exc),
                })

    return results


def render_css() -> None:
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Instrument+Sans:wght@400;500;600;700&family=Newsreader:opsz,wght@6..72,500;6..72,650&display=swap');

        :root {
            --bg: #000000;
            --panel: #0d0f13;
            --panel-2: #151820;
            --ink: #f5f7fb;
            --muted: #8d96a6;
            --line: #222733;
            --accent: #54d6b6;
            --accent-2: #a5b4fc;
            --warn: #f6c267;
            --bad: #fb8b8b;
            --shadow: 0 18px 55px rgba(0, 0, 0, 0.55);
        }

        html, body, [class*="css"], .stApp {
            font-family: 'Instrument Sans', system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            color: var(--ink);
        }

        .stApp {
            background:
                linear-gradient(180deg, rgba(0,0,0,0.98), rgba(0,0,0,1)),
                radial-gradient(circle at top left, rgba(84,214,182,0.08), transparent 30%),
                var(--bg);
        }

        .block-container {
            max-width: 1160px;
            padding-top: 2rem;
            padding-bottom: 3rem;
        }

        h1, h2, h3, .hero-title {
            font-family: 'Newsreader', Georgia, serif !important;
            letter-spacing: 0 !important;
            color: var(--ink);
        }

        h2, h3 { font-weight: 650 !important; }

        .hero {
            padding: 30px 0 20px;
            max-width: 860px;
        }

        .eyebrow {
            color: var(--accent);
            font-size: 0.76rem;
            font-weight: 700;
            letter-spacing: 0.12em;
            text-transform: uppercase;
            margin-bottom: 8px;
        }

        .hero-title {
            font-size: clamp(2.4rem, 6vw, 5rem);
            line-height: 0.95;
            font-weight: 650;
            margin: 0 0 14px;
        }

        .hero-copy {
            max-width: 680px;
            color: var(--muted);
            font-size: 1.02rem;
            line-height: 1.7;
            margin: 0;
        }

        [data-testid="stSidebar"] {
            background: #050506;
            border-right: 1px solid var(--line);
        }

        [data-testid="stSidebar"] * {
            color: #eef2f7;
        }

        [data-testid="stSidebar"] p,
        [data-testid="stSidebar"] .stCaptionContainer {
            color: #98a2b3 !important;
        }

        [data-testid="stMetric"] {
            background: var(--panel);
            border: 1px solid var(--line);
            border-radius: 12px;
            padding: 15px 16px;
            box-shadow: var(--shadow);
        }

        [data-testid="stMetricLabel"] p {
            color: var(--muted) !important;
            font-size: 0.78rem !important;
        }

        .joy-card {
            background: var(--panel);
            border: 1px solid var(--line);
            border-radius: 12px;
            padding: 16px;
            box-shadow: var(--shadow);
        }

        .muted { color: var(--muted); }

        .small-label {
            color: var(--muted);
            font-size: 0.78rem;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            margin-bottom: 0.35rem;
        }
        .success-pill, .warn-pill, .bad-pill {
            display: inline-block;
            padding: 3px 8px;
            border-radius: 999px;
            font-size: 0.78rem;
            border: 1px solid var(--line);
        }
        .success-pill { color: var(--accent); }
        .warn-pill { color: var(--warn); }
        .bad-pill { color: var(--bad); }

        [data-baseweb="tab-list"] {
            gap: 8px;
            border-bottom: 1px solid var(--line);
        }

        [data-baseweb="tab"] {
            font-weight: 650;
            color: var(--muted);
            padding-left: 4px;
            padding-right: 18px;
        }

        [aria-selected="true"] {
            color: var(--ink) !important;
        }

        textarea, input {
            border-radius: 10px !important;
            border-color: var(--line) !important;
            background: #0b0d11 !important;
            color: var(--ink) !important;
        }

        textarea:focus, input:focus {
            border-color: var(--accent) !important;
            box-shadow: 0 0 0 3px rgba(84,214,182,0.14) !important;
        }

        [data-testid="stFileUploaderDropzone"] {
            background: #0b0d11;
            border: 1px dashed #303746;
            border-radius: 12px;
        }

        [data-testid="stFileUploaderDropzone"] * {
            color: var(--muted) !important;
        }

        .stButton button,
        .stDownloadButton button,
        [data-testid="stFormSubmitButton"] button {
            border-radius: 10px !important;
            font-weight: 700 !important;
            letter-spacing: 0 !important;
        }

        .stButton button[kind="primary"],
        .stDownloadButton button[kind="primary"] {
            background: var(--ink) !important;
            border-color: var(--ink) !important;
            color: #000 !important;
        }

        div[data-testid="stDataFrame"] {
            border: 1px solid var(--line);
            border-radius: 12px;
            overflow: hidden;
            box-shadow: var(--shadow);
        }

        .stAlert {
            border-radius: 12px;
        }
        
        button[kind="secondary"] {
            background: transparent !important;
            border: 1px solid #222733 !important;
            color: #98a2b3 !important;
            transition: all 0.18s ease;
            font-weight: 600 !important;
        }

        button[kind="secondary"]:hover {
            border-color: #54d6b6 !important;
            color: white !important;
            background: rgba(84,214,182,0.06) !important;
        }

        </style>
        """,
        unsafe_allow_html=True,
    )


def show_results_summary(df: pd.DataFrame) -> None:
    if df.empty:
        return

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Screened", len(df))
    c2.metric("Strong Fit", int((df["Verdict"] == "Strong Fit").sum()))
    c3.metric("Good Fit", int((df["Verdict"] == "Good Fit").sum()))
    c4.metric("Average Score", round(float(df["Final Score"].mean()), 1))

    display_cols = [
        "Send", "Name", "Email", "Phone", "Experience", "Final Score", "Verdict",
        "Matched Keywords", "Missing Keywords", "Reason", "Source File",
    ]
    st.dataframe(df[display_cols], use_container_width=True, hide_index=True)

    st.download_button(
        "Download screening CSV",
        df.to_csv(index=False).encode("utf-8"),
        "joy_screening_results.csv",
        "text/csv",
        use_container_width=False,
    )

def process_resume_worker(
    upload,
    jd_text,
    role,
    keywords,
    min_exp,
    api_key,
    model,
):

    if not upload:
        return None

    text, error = read_uploaded_file(
        upload.name,
        upload.getvalue(),
    )

    if error:
        return {
            "Send": False,
            "Name": upload.name,
            "Email": "",
            "Phone": "",
            "Experience": 0.0,
            "Keyword Score": 0,
            "Final Score": 0.0,
            "Verdict": "Low Fit",
            "Matched Keywords": "",
            "Missing Keywords": ", ".join(keywords[:10]),
            "Skills": "",
            "Reason": error,
            "Source File": upload.name,
            "AI Used": False,
        }

    if not text.strip():
        return {
            "Send": False,
            "Name": upload.name,
            "Email": "",
            "Phone": "",
            "Experience": 0.0,
            "Keyword Score": 0,
            "Final Score": 0.0,
            "Verdict": "Low Fit",
            "Matched Keywords": "",
            "Missing Keywords": ", ".join(keywords[:10]),
            "Skills": "",
            "Reason": "No readable text found.",
            "Source File": upload.name,
            "AI Used": False,
        }

    return score_resume(
        jd_text=jd_text,
        role=role,
        resume_text=text,
        filename=upload.name,
        keywords=keywords,
        min_exp=min_exp,
        api_key=api_key,
        model=model,
    )



def run_screening(
    uploads,
    jd_text: str,
    role_input: str,
    extra_keywords: str,
    api_key: str,
    model: str,
    user_key: str,
) -> tuple[pd.DataFrame, list[str]]:
    errors = []
    role = detect_role_title(jd_text, role_input, api_key, model)
    keywords = extract_keywords(f"{role}\n{jd_text}", extra_keywords)
    min_exp = parse_min_experience(jd_text)
    rows = []

    progress = st.progress(0)
    status = st.empty()

    cpu_count = os.cpu_count() or 4

    max_workers = min(
        4,
        cpu_count,
        max(1, len(uploads))
    )

    with ThreadPoolExecutor(max_workers=max_workers) as executor:

        futures = {
            executor.submit(
                process_resume_worker,
                upload,
                jd_text,
                role,
                keywords,
                min_exp,
                api_key,
                model,
            ): upload
            for upload in uploads
        }

        completed = 0

        for future in as_completed(futures):

            upload = futures[future]

            try:
                result = future.result()
                rows.append(result)

            except Exception as exc:
                errors.append(f"{upload.name}: {exc}")

            completed += 1

            status.write(
                f"Processed {completed}/{len(uploads)} resumes"
            )

            progress.progress(completed / len(uploads))

    progress.empty()
    status.empty()

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values("Final Score", ascending=False).reset_index(drop=True)
        df.insert(0, "Rank", range(1, len(df) + 1))

    st.session_state.last_role = role
    st.session_state.last_jd = jd_text
    st.session_state.last_keywords = keywords
    st.session_state.results_df = df
    save_history(df, role, user_key)
    return df, errors


st.set_page_config(page_title=f"{APP_NAME} AI Recruiter", page_icon="J", layout="wide")
render_css()
init_state()

if not st.session_state.gmail_authenticated:
    st.markdown(
        """
        <section class="hero">
            <div class="eyebrow">Joy AI Recruiter</div>
            <h1 class="hero-title">Screen once. Ask once.</h1>
            <p class="hero-copy">
                Sign in with the Gmail account that should send candidate emails.
                Credentials stay in this Streamlit session and are never written to GitHub.
            </p>
        </section>
        """,
        unsafe_allow_html=True,
    )

    with st.form("gmail_login"):
        login_email = st.text_input("Gmail address", placeholder="you@gmail.com")
        login_password = st.text_input("Gmail App Password", type="password", placeholder="16-character app password")
        login_name = st.text_input("Your name", placeholder="Auto-filled from email if left blank")
        login_company = st.text_input("Company", value=st.session_state.company_name or DEFAULT_COMPANY)
        submitted = st.form_submit_button("Start screening", type="primary", use_container_width=True)

    st.caption("Use a Gmail App Password from Google Account > Security > 2-Step Verification > App Passwords.")

    if submitted:
        if "@" not in login_email:
            st.error("Enter a valid Gmail address.")
        elif len(login_password.replace(" ", "").strip()) < 16:
            st.error("Enter your Gmail App Password.")
        else:
            login_user(login_email, login_password, login_name, login_company)
            st.rerun()

    st.stop()

with st.sidebar:
    st.title("Joy")
    st.caption("Screen. Select. Send.")

    st.session_state.sender_name = st.text_input(
        "Sender name",
        value=st.session_state.sender_name,
        placeholder="Your name",
    )
    st.session_state.company_name = st.text_input(
        "Company",
        value=st.session_state.company_name,
        placeholder=DEFAULT_COMPANY,
    )

    st.divider()
    st.caption(f"Signed in as {mask_email(st.session_state.sender_email)}")
    if st.button("Change Gmail login", use_container_width=True):
        logout_user()

    st.divider()
    openai_api_key = get_secret("OPENAI_API_KEY")
    openai_model = get_secret("OPENAI_MODEL", "gpt-4o-mini")
    ai_status = "AI scoring enabled" if openai_api_key else "Heuristic scoring active"
    st.caption(ai_status)

    st.divider()
    user_key = st.session_state.sender_email or "local"
    if st.button("Clear current results", use_container_width=True):
        st.session_state.results_df = pd.DataFrame()
        st.session_state.email_results = []
        st.rerun()


st.markdown(
    """
    <section class="hero">
        <div class="eyebrow">Joy AI Recruiter</div>
        <h1 class="hero-title">Screen once. Ask once.</h1>
        <p class="hero-copy">
            Rank resumes against one role, then send a precise email that collects CTC,
            notice period, location, availability, and fit details before any call.
        </p>
    </section>
    """,
    unsafe_allow_html=True,
)

screen_tab, email_tab, history_tab = st.tabs(["Screen", "Email", "History"])

with screen_tab:

    title_col, button_col = st.columns([8.5, 1.5], vertical_alignment="center")

    with title_col:
        st.subheader("Job")

    with button_col:
        st.markdown("<div style='height: 30px'></div>", unsafe_allow_html=True)

        new_search = st.button(
            "New",
            key="new_search_btn",
            use_container_width=True,
        )

    if new_search:

        reset_screening_session()

        if not st.session_state.get("_reset_done"):
            st.rerun()

        st.session_state["_reset_done"] = False

    jd_upload = st.file_uploader(
        "Upload JD",
        type=["pdf", "docx", "txt"],
        key=f"jd_upload_{st.session_state.upload_session}",
    )

    typed_jd_text = st.text_area(
        "Or paste JD",
        height=190,
        placeholder="Paste the job description or role requirements here. Joy will detect the title automatically.",
        key="typed_jd_text",
    )

    jd_text = typed_jd_text
    if jd_upload:
        uploaded_jd_text, jd_error = read_uploaded_file(
            jd_upload.name,
            jd_upload.getvalue(),
        )
        if jd_error:
            st.warning(f"JD upload: {jd_error}")
        if uploaded_jd_text.strip():
            jd_text = uploaded_jd_text
            st.caption(f"Using uploaded JD: {jd_upload.name}")

    with st.expander("Optional screening controls", expanded=False):
        role_input = st.text_input(
            "Role title override",
            placeholder="Leave blank. Joy will detect it from the JD.",
            key="role_input",
        )
        extra_keywords = st.text_input(
            "Must-have keywords",
            placeholder="HPLC, distributor management, SAP",
            key="extra_keywords",
        )

    detected_preview = extract_role_from_jd(jd_text, role_input) if (jd_text.strip() or role_input.strip()) else ""
    if detected_preview and detected_preview != "Open Role":
        st.caption(f"Detected role title: {detected_preview}")

    uploads = st.file_uploader(
        "Upload resumes",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key=f"resume_uploads_{st.session_state.upload_session}",
    )

    run_col, _ = st.columns([1, 4])
    with run_col:
        run_clicked = st.button("Screen resumes", type="primary", use_container_width=True)

    if run_clicked:
        if not uploads:
            st.error("Upload at least one resume.")
        elif not role_input.strip() and not jd_text.strip():
            st.error("Upload or paste a JD, or add a role override in Optional screening controls.")
        else:
            with st.spinner("Screening resumes..."):
                results, read_errors = run_screening(
                    uploads=uploads,
                    jd_text=jd_text,
                    role_input=role_input,
                    extra_keywords=extra_keywords,
                    api_key=openai_api_key,
                    model=openai_model,
                    user_key=user_key,
                )
            st.success(f"Screened {len(results)} resume(s) for {st.session_state.last_role}.")
            for error in read_errors:
                st.warning(error)

    if not st.session_state.results_df.empty:
        st.divider()
        st.subheader(f"Results: {st.session_state.last_role}")
        show_results_summary(st.session_state.results_df)


with email_tab:
    st.subheader("Outreach")

    if st.session_state.results_df.empty:
        st.info("Run a screening first.")
    else:
        editable = st.session_state.results_df.copy()
        editable["Send"] = editable["Send"].astype(bool)

        edited = st.data_editor(
            editable,
            use_container_width=True,
            hide_index=True,
            num_rows="fixed",
            disabled=[
                "Rank", "Phone", "Experience", "Keyword Score", "Final Score",
                "Verdict", "Matched Keywords", "Missing Keywords", "Skills",
                "Reason", "Source File", "AI Used",
            ],
            column_config={
                "Send": st.column_config.CheckboxColumn("Send"),
                "Email": st.column_config.TextColumn("Email"),
            },
            key="email_editor",
        )

        selected = edited[edited["Send"] == True].copy()
        missing_email = selected[~selected["Email"].astype(str).str.contains("@", na=False)]

        st.session_state.questions_text = st.text_area(
            "Questions to collect",
            value=st.session_state.questions_text,
            height=180,
        )
        extra_note = st.text_area(
            "Extra note",
            placeholder="Example: This is a hybrid Pune role with a 25-30 LPA budget.",
            height=90,
        )

        subject = st.text_input(
            "Subject",
            value=f"Details required for {st.session_state.last_role} opportunity",
        )

        questions = questions_from_text(st.session_state.questions_text)
        
        if not selected.empty:
            preview_body = build_email_body(
                selected.iloc[0],
                st.session_state.last_role,
                st.session_state.sender_name,
                st.session_state.company_name,
                questions,
                extra_note,
            )


        if not selected.empty:

            with st.expander(
                f"Preview: {selected.iloc[0]['Name']}",
                expanded=True
            ):

                edited_preview_body = st.text_area(
                    "Edit email before sending",
                    value=preview_body,
                    height=380,
                    key="edited_email_preview",
                )

                st.caption("Use {first_name} anywhere for automatic personalization.")

                st.caption(
                    "Variables supported: "
                    "{first_name}, {full_name}, {role}, "
                    "{experience}, {score}, {verdict}"
                )

        c1, c2, c3 = st.columns([1.3, 1.5, 3])
        with c1:
            confirm = st.checkbox("Recipient list reviewed")
        with c2:
            send_clicked = st.button(
                f"Send {len(selected)} email(s)",
                type="primary",
                disabled=selected.empty or not confirm,
                use_container_width=True,
            )

        if not missing_email.empty:
            st.warning("Add valid email addresses before sending: " + ", ".join(missing_email["Name"].astype(str).tolist()))

        if send_clicked:
            if not st.session_state.sender_email or not st.session_state.sender_password:
                st.error("Your Gmail session expired. Sign in again to send emails.")

            elif not st.session_state.sender_name:
                st.error("Add sender name in the sidebar.")

            elif not missing_email.empty:
                st.error("Fix missing candidate email addresses first.")

            else:
                custom_email_body = st.session_state.get("edited_email_preview", "").strip()

                with st.spinner("Sending emails..."):
                    st.session_state.email_results = send_bulk_emails(
                        selected_df=selected,
                        role=st.session_state.last_role,
                        sender_email=st.session_state.sender_email,
                        sender_password=st.session_state.sender_password,
                        sender_name=st.session_state.sender_name,
                        company_name=st.session_state.company_name,
                        subject=subject,
                        questions=questions,
                        extra_note=extra_note,
                        custom_body=custom_email_body,
                    )

                sent_count = sum(
                    1 for item in st.session_state.email_results if item["Success"]
                )

                st.success(
                    f"Sent {sent_count} of {len(st.session_state.email_results)} email(s)."
                )

            if st.session_state.email_results:
                st.dataframe(pd.DataFrame(st.session_state.email_results), use_container_width=True, hide_index=True)


with history_tab:
    st.subheader("History")
    hist = load_history(user_key)

    if hist.empty:
        st.info("No saved screenings yet.")
    else:
        c1, c2, c3 = st.columns(3)
        c1.metric("Candidates", len(hist))
        c2.metric("Strong Fit", int((hist["Verdict"] == "Strong Fit").sum()) if "Verdict" in hist.columns else 0)
        c3.metric("Roles", hist["Role"].nunique() if "Role" in hist.columns else 0)

        if "Role" in hist.columns:
            roles = ["all"] + sorted(hist["Role"].dropna().unique().tolist())
            selected_role = st.selectbox("Role filter", roles)
            shown = hist if selected_role == "all" else hist[hist["Role"] == selected_role]
            delete_col1, delete_col2 = st.columns(2)

            with delete_col1:

                if selected_role != "all":

                    if st.button(
                        f"Delete {selected_role} history",
                        use_container_width=True,
                    ):

                        clear_role_history(user_key, selected_role)

                        st.success(f"Deleted history for {selected_role}")

                        st.rerun()

            with delete_col2:

                if st.button(
                    "Delete all history",
                    use_container_width=True,
                ):

                    clear_history(user_key)

                    st.success("All history deleted")

                    st.rerun()
        else:
            shown = hist

        history_editable = shown.copy()
        # limit huge history rendering slowdown
        history_editable = history_editable.tail(300)

        if "Send" not in history_editable.columns:
            history_editable.insert(0, "Send", False)

        history_edited = st.data_editor(
            history_editable,
            height=500,
            use_container_width=True,
            hide_index=True,
            num_rows="fixed",
            key="history_editor",
            column_config={
                "Send": st.column_config.CheckboxColumn("Send"),
                "Email": st.column_config.TextColumn("Email"),
            },
        )

        selected_history = history_edited[history_edited["Send"] == True].copy()

        if not selected_history.empty:

            st.divider()
            st.subheader("Send email from history")

            history_role = selected_history.iloc[0].get(
                "Role",
                st.session_state.last_role or "the role",
            )

            history_subject = st.text_input(
                "Subject",
                value=f"Details required for {history_role}",
                key="history_subject",
            )

            history_questions = st.text_area(
                "Questions to collect",
                value=st.session_state.questions_text,
                height=180,
                key="history_questions",
            )

            history_note = st.text_area(
                "Extra note",
                placeholder="Optional context for candidates",
                height=100,
                key="history_note",
            )

            parsed_questions = questions_from_text(history_questions)

            preview_body = build_email_body(
                selected_history.iloc[0],
                selected_history.iloc[0].get("Role", st.session_state.last_role),
                st.session_state.sender_name,
                st.session_state.company_name,
                parsed_questions,
                history_note,
            )

            edited_history_body = st.text_area(
                "Edit email before sending",
                value=preview_body,
                height=380,
                key="history_email_preview",
            )

            history_confirm = st.checkbox(
                "History recipient list reviewed",
                key="history_confirm",
            )

            send_history = st.button(
                f"Send {len(selected_history)} email(s)",
                type="primary",
                disabled=not history_confirm,
                key="send_history_btn",
            )

            if send_history:

                custom_body = st.session_state.get(
                    "history_email_preview",
                    "",
                ).strip()

                with st.spinner("Sending emails from history..."):

                    history_results = send_bulk_emails(
                        selected_df=selected_history,
                        role=selected_history.iloc[0].get(
                            "Role",
                            st.session_state.last_role,
                        ),
                        sender_email=st.session_state.sender_email,
                        sender_password=st.session_state.sender_password,
                        sender_name=st.session_state.sender_name,
                        company_name=st.session_state.company_name,
                        subject=history_subject,
                        questions=parsed_questions,
                        extra_note=history_note,
                        custom_body=custom_body,
                    )

                sent_count = sum(
                    1 for item in history_results if item["Success"]
                )

                st.success(
                    f"Sent {sent_count} of {len(history_results)} email(s)."
                )

                st.dataframe(
                    pd.DataFrame(history_results),
                    use_container_width=True,
                    hide_index=True,
                )
