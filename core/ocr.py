import re
import threading
from io import BytesIO

import streamlit as st

try:
    import pdfplumber
except Exception:
    pdfplumber = None
try:
    from docx import Document
except Exception:
    Document = None
try:
    import pytesseract
except Exception:
    pytesseract = None
try:
    from pdf2image import convert_from_bytes
except Exception:
    convert_from_bytes = None


# ---------------------------------------------------------------------------
# Tesseract path — only set on Windows. On Streamlit Cloud / Linux, tesseract
# is already on PATH via packages.txt, and forcing a Windows path there is
# dead code that does nothing (silently swallowed by the try/except) but
# confuses anyone reading this file wondering which OS it's meant for.
# ---------------------------------------------------------------------------
import platform

if pytesseract is not None and platform.system() == "Windows":
    default_tesseract = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    try:
        pytesseract.pytesseract.tesseract_cmd = default_tesseract
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Per-file locks instead of one global lock.
# The old global `ocr_lock` serialized OCR across ALL users and ALL files —
# if two people screened scanned resumes at the same moment, one blocked
# the other for the full OCR duration (which is already your slowest path:
# rasterizing + Tesseract). convert_from_bytes and pytesseract are safe to
# run concurrently on independent byte inputs; nothing here needs a shared
# lock. Removed entirely — @st.cache_data already prevents duplicate work
# on the *same* file.
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# GARBAGE / LOW-QUALITY TEXT DETECTION
#
# The old check (`alpha_ratio < 0.25`) only catches PDFs that extracted to
# near-nothing. It misses the more common and more damaging failure mode:
# text that extracted with a HIGH alpha ratio but is functionally useless —
# letter-spaced garbage from certain font encodings ("J oh n D oe"), or
# extraction where words got glued together with no whitespace at all
# ("JohnDoeSoftwareEngineer"). Both pass the old check and get silently
# scored on garbage, with no OCR fallback ever triggered.
# ---------------------------------------------------------------------------
def _looks_letter_spaced(text: str) -> bool:
    """True if a large share of 'words' are 1-2 character alphabetic
    fragments — the signature of broken font-encoding extraction, whether
    it's fully spaced ('J o h n D o e') or partially fragmented
    ('J oh n D oe Sof twa re'), which is the more common real-world case.
    A pure single-letter check misses partial fragmentation, which is why
    this counts 1-2 char tokens rather than only 1 char tokens."""
    if not text:
        return False
    tokens = text.split()
    if len(tokens) < 20:
        return False
    short_fragments = sum(1 for t in tokens if len(t) <= 2 and t.isalpha())
    return (short_fragments / len(tokens)) > 0.40


def _looks_glued_together(text: str) -> bool:
    """True if the text has almost no whitespace relative to its length —
    a sign that words/lines were merged with no separators during extraction
    (common with certain multi-column resume templates)."""
    if not text or len(text) < 200:
        return False
    space_ratio = text.count(" ") / len(text)
    return space_ratio < 0.06


def _extraction_quality_ok(text: str) -> bool:
    """Single gate deciding whether extracted text is usable, or whether
    we should fall back to OCR. Combines density, letter-spacing, and
    glued-text checks instead of relying on alpha_ratio alone."""
    if not text or not text.strip():
        return False

    alpha_ratio = sum(c.isalpha() for c in text) / max(len(text), 1)
    if alpha_ratio < 0.25:
        return False
    if _looks_letter_spaced(text):
        return False
    if _looks_glued_together(text):
        return False
    return True


def repair_letter_spaced_text(text: str) -> str:
    """
    Collapse ' J o h n   D o e ' style extraction artifacts back into
    normal words, applied to the FULL document text at the extraction
    layer — not just to name candidates. This means keyword extraction,
    skill matching, and experience parsing all benefit, not only the
    name detector.
    """
    if not text:
        return text

    out_lines = []
    for line in text.splitlines():
        tokens = line.split(" ")
        if len(tokens) < 4:
            out_lines.append(line)
            continue

        single_letter_ratio = sum(
            1 for t in tokens if len(t) == 1 and t.isalpha()
        ) / max(len(tokens), 1)

        if single_letter_ratio < 0.5:
            out_lines.append(line)
            continue

        words, current = [], []
        for tok in tokens:
            if tok == "":
                if current:
                    words.append("".join(current))
                    current = []
                continue
            if len(tok) == 1 and tok.isalpha():
                current.append(tok)
            else:
                if current:
                    words.append("".join(current))
                    current = []
                words.append(tok)
        if current:
            words.append("".join(current))

        out_lines.append(" ".join(w for w in words if w))

    return "\n".join(out_lines)


def ocr_pdf(data: bytes, max_pages: int = 8) -> str:
    """OCR fallback. max_pages now matches the text-extraction page cap
    below (was capped at 5 while the text path allowed 8 — inconsistent,
    and could silently drop OCR-worthy content on pages 6-8)."""
    if pytesseract is None or convert_from_bytes is None:
        return ""
    try:
        images = convert_from_bytes(data, dpi=250, first_page=1, last_page=max_pages)
        text_parts = []
        for image in images:
            text_parts.append(
                pytesseract.image_to_string(image, config="--psm 3 --oem 3")
            )
        return "\n".join(text_parts).strip()
    except Exception:
        return ""


MAX_PAGES = 8


@st.cache_data(show_spinner=False, max_entries=250)
def read_uploaded_file(file_name: str, data: bytes) -> tuple[str, str]:
    name = file_name.lower()
    try:
        if name.endswith(".pdf"):
            if pdfplumber is None:
                return "", "pdfplumber is not installed."

            with pdfplumber.open(BytesIO(data)) as pdf:
                total_pages = len(pdf.pages)
                pages = pdf.pages[:MAX_PAGES]
                text = "\n".join(page.extract_text() or "" for page in pages).strip()

            truncated_note = ""
            if total_pages > MAX_PAGES:
                truncated_note = (
                    f" (Note: PDF has {total_pages} pages; only first {MAX_PAGES} were read.)"
                )

            if not _extraction_quality_ok(text):
                ocr_text = ocr_pdf(data, max_pages=MAX_PAGES)
                # Prefer OCR text if it's meaningfully longer AND passes
                # the same quality gate — don't swap in OCR output that's
                # itself garbage just because it happens to be longer.
                if len(ocr_text) > len(text) and _extraction_quality_ok(ocr_text):
                    text = ocr_text
                elif not text.strip() and ocr_text.strip():
                    # Original extraction was empty; OCR text is our only
                    # option even if imperfect — better than nothing.
                    text = ocr_text

            text = repair_letter_spaced_text(text)

            if not text.strip():
                return "", "No readable text found (scanned/image-only PDF, OCR unavailable or failed)."

            return text.strip(), truncated_note.strip()

        if name.endswith(".docx"):
            if Document is None:
                return "", "python-docx is not installed."
            doc = Document(BytesIO(data))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(
                        cell.text for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
            text = "\n".join(text_parts).strip()
            text = repair_letter_spaced_text(text)
            if not text:
                return "", "DOCX opened but no readable text found."
            return text, ""

        if name.endswith(".txt"):
            text = data.decode("utf-8", errors="ignore").strip()
            return repair_letter_spaced_text(text), ""

        return "", "Unsupported file type."
    except Exception as exc:
        return "", f"Could not read file: {exc}"
